"""Document generation router for AI Act compliance documents."""
from typing import List, Optional, Dict, Any
from uuid import UUID
from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException, BackgroundTasks
from sqlalchemy.orm import Session
from pydantic import BaseModel
from jinja2 import Environment, FileSystemLoader
import tempfile
import os

from app.db.session import get_db
from app.core.auth import get_current_user

router = APIRouter(prefix="/documents", tags=["documents"])

# Jinja2 setup
template_dir = os.path.join(os.path.dirname(__file__), "..", "templates", "documents")
jinja_env = Environment(loader=FileSystemLoader(template_dir))

# Document types configuration
DOCUMENT_TYPES = {
    "ai_policy": {
        "title": "Política de Uso de IA",
        "description": "Documento maestro de gobernanza de IA",
        "template": "ai_policy.html.j2",
        "requires_use_case": False,
        "obligatory_for": ["high", "limited", "minimal"]
    },
    "employee_notice": {
        "title": "Aviso de Uso de IA a Empleados",
        "description": "Información a empleados sobre sistemas de IA",
        "template": "employee_notice.html.j2",
        "requires_use_case": False,
        "obligatory_for": ["high"]
    },
    "systems_register": {
        "title": "Registro de Sistemas de IA",
        "description": "Inventario formal de todos los sistemas de IA",
        "template": "systems_register.html.j2",
        "requires_use_case": False,
        "obligatory_for": ["high", "limited", "minimal"]
    },
    "fria": {
        "title": "Evaluación de Impacto en Derechos Fundamentales (FRIA)",
        "description": "Evaluación previa al despliegue de sistemas de alto riesgo",
        "template": "fria.html.j2",
        "requires_use_case": True,
        "obligatory_for": ["high"]
    },
    "candidate_notice": {
        "title": "Aviso de IA en Procesos de Selección",
        "description": "Información a candidatos sobre uso de IA en RRHH",
        "template": "candidate_notice.html.j2",
        "requires_use_case": True,
        "obligatory_for": ["high"]
    }
}

# Schemas
class GenerateDocumentRequest(BaseModel):
    type: str
    use_case_id: Optional[UUID] = None

class DocumentResponse(BaseModel):
    id: UUID
    title: str
    type: str
    pdf_url: Optional[str] = None
    docx_url: Optional[str] = None
    generated_at: datetime
    
    class Config:
        from_attributes = True

class DocumentListResponse(BaseModel):
    documents: List[DocumentResponse]


def get_document_config(doc_type: str) -> Dict[str, Any]:
    """Get document type configuration."""
    if doc_type not in DOCUMENT_TYPES:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid document type. Must be one of: {list(DOCUMENT_TYPES.keys())}"
        )
    return DOCUMENT_TYPES[doc_type]


def render_html_template(template_name: str, data: Dict[str, Any]) -> str:
    """Render HTML template with Jinja2."""
    template = jinja_env.get_template(template_name)
    return template.render(**data)


def generate_pdf_from_html(html_content: str) -> bytes:
    """Generate PDF from HTML content.
    
    Note: This is a placeholder. In production, use weasyprint or similar.
    """
    try:
        import weasyprint
        return weasyprint.HTML(string=html_content).write_pdf()
    except ImportError:
        # Fallback: return HTML as bytes for now
        return html_content.encode('utf-8')


def generate_docx_from_data(data: Dict[str, Any], doc_type: str) -> bytes:
    """Generate DOCX from data.
    
    Note: This is a placeholder. In production, use python-docx properly.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Add title
        title = doc.add_heading(data.get('organization', {}).get('name', 'Documento'), 0)
        
        # Add document type title
        config = DOCUMENT_TYPES.get(doc_type, {})
        doc.add_heading(config.get('title', 'Documento'), level=1)
        
        # Add metadata
        doc.add_paragraph(f"Fecha: {data.get('generated_date', datetime.now().strftime('%Y-%m-%d'))}")
        doc.add_paragraph()
        
        # Add use cases if present
        use_cases = data.get('use_cases', [])
        if use_cases:
            doc.add_heading('Sistemas de IA', level=2)
            for uc in use_cases:
                p = doc.add_paragraph()
                p.add_run(f"• {uc.get('name', 'Sin nombre')}").bold = True
                doc.add_paragraph(f"  Herramienta: {uc.get('tool_name', 'N/A')}")
                doc.add_paragraph(f"  Proveedor: {uc.get('tool_provider', 'N/A')}")
                doc.add_paragraph(f"  Clasificación: {uc.get('risk_classification', 'N/A')}")
        
        # Save to bytes
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc.save(tmp.name)
            with open(tmp.name, 'rb') as f:
                content = f.read()
            os.unlink(tmp.name)
            return content
    except ImportError:
        # Fallback: return simple text as bytes
        return f"Documento DOCX - {doc_type}".encode('utf-8')


@router.post("/generate", response_model=DocumentResponse)
def generate_document(
    request: GenerateDocumentRequest,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """Generate a compliance document (PDF + DOCX).
    
    Requires Pro or Agency plan.
    """
    from app.models.subscription import Subscription
    
    # Check subscription plan
    subscription = db.query(Subscription).filter(
        Subscription.organization_id == current_user.organization_id,
        Subscription.status == "active"
    ).first()
    
    allowed_plans = ["pro", "agency"]
    if not subscription or subscription.plan not in allowed_plans:
        raise HTTPException(
            status_code=403,
            detail="Document generation requires Pro or Agency plan"
        )
    
    # Get document config
    config = get_document_config(request.type)
    
    # Validate use_case_id if required
    if config["requires_use_case"] and not request.use_case_id:
        raise HTTPException(
            status_code=400,
            detail=f"Document type '{request.type}' requires use_case_id"
        )
    
    # Fetch organization data
    from app.models.organization import Organization
    organization = db.query(Organization).filter(
        Organization.id == current_user.organization_id
    ).first()
    
    if not organization:
        raise HTTPException(status_code=404, detail="Organization not found")
    
    # Fetch use cases
    from app.models.use_case import UseCase
    use_cases_query = db.query(UseCase).filter(
        UseCase.organization_id == current_user.organization_id,
        UseCase.deleted_at.is_(None)
    )
    
    if request.use_case_id:
        use_case = use_cases_query.filter(UseCase.id == request.use_case_id).first()
        if not use_case:
            raise HTTPException(status_code=404, detail="Use case not found")
        use_cases = [use_case]
    else:
        use_cases = use_cases_query.all()
    
    # Filter use cases based on document type
    if request.type == "employee_notice":
        # Filter for HR department use cases
        use_cases = [uc for uc in use_cases if uc.department and uc.department.lower() in ["rrhh", "hr", "recursos humanos", "human resources"]]
    elif request.type == "candidate_notice":
        # Filter for HR department high-risk use cases
        use_cases = [uc for uc in use_cases if uc.department and uc.department.lower() in ["rrhh", "hr", "recursos humanos", "human resources"] and uc.ai_act_level == "high"]
    elif request.type == "fria":
        # For FRIA, only include the specific high-risk use case
        use_cases = [uc for uc in use_cases if uc.ai_act_level == "high"]
    
    # Prepare template data
    template_data = {
        "organization": {
            "name": organization.name,
            "sector": organization.sector or "No especificado"
        },
        "use_cases": [
            {
                "name": uc.name,
                "description": uc.description,
                "tool_name": uc.tool_name,
                "tool_provider": uc.tool_provider,
                "department": uc.department,
                "risk_classification": uc.ai_act_level,
                "obligations_summary": uc.obligations_summary or [],
                "applicable_articles": uc.applicable_articles or "Art. 6, 52",
                "questionnaire_answers": uc.questionnaire_answers or []
            }
            for uc in use_cases
        ],
        "generated_date": datetime.now().strftime("%d/%m/%Y"),
        "use_case": use_cases[0] if use_cases and request.type == "fria" else None,
        "questionnaire_answers": use_cases[0].questionnaire_answers if use_cases and request.type == "fria" else []
    }
    
    # Render HTML template
    html_content = render_html_template(config["template"], template_data)
    
    # Generate PDF
    pdf_content = generate_pdf_from_html(html_content)
    
    # Generate DOCX
    docx_content = generate_docx_from_data(template_data, request.type)
    
    # Upload to Supabase Storage
    from app.core.supabase import supabase_client
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    org_id = str(current_user.organization_id)
    
    pdf_path = f"{org_id}/{request.type}_{timestamp}.pdf"
    docx_path = f"{org_id}/{request.type}_{timestamp}.docx"
    
    try:
        # Upload PDF
        supabase_client.storage.from_("documents").upload(
            pdf_path,
            pdf_content,
            {"content-type": "application/pdf"}
        )
        
        # Upload DOCX
        supabase_client.storage.from_("documents").upload(
            docx_path,
            docx_content,
            {"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
        )
        
        # Generate signed URLs (1 hour expiry)
        pdf_url = supabase_client.storage.from_("documents").create_signed_url(pdf_path, 3600)
        docx_url = supabase_client.storage.from_("documents").create_signed_url(docx_path, 3600)
        
    except Exception as e:
        # If storage fails, still return the document record
        pdf_url = None
        docx_url = None
    
    # Create document record in database
    from app.models.document import Document
    
    document = Document(
        organization_id=current_user.organization_id,
        type=request.type,
        title=config["title"],
        pdf_path=pdf_path if pdf_url else None,
        docx_path=docx_path if docx_url else None,
        pdf_url=pdf_url.get("signedURL") if pdf_url else None,
        docx_url=docx_url.get("signedURL") if docx_url else None,
        generated_at=datetime.utcnow(),
        use_case_id=request.use_case_id
    )
    
    db.add(document)
    db.commit()
    db.refresh(document)
    
    return DocumentResponse(
        id=document.id,
        title=document.title,
        type=document.type,
        pdf_url=document.pdf_url,
        docx_url=document.docx_url,
        generated_at=document.generated_at
    )


@router.get("", response_model=DocumentListResponse)
def list_documents(
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """List all generated documents for the organization."""
    from app.models.document import Document
    
    documents = db.query(Document).filter(
        Document.organization_id == current_user.organization_id
    ).order_by(Document.generated_at.desc()).all()
    
    return DocumentListResponse(
        documents=[
            DocumentResponse(
                id=doc.id,
                title=doc.title,
                type=doc.type,
                pdf_url=doc.pdf_url,
                docx_url=doc.docx_url,
                generated_at=doc.generated_at
            )
            for doc in documents
        ]
    )


@router.get("/{document_id}/download")
def download_document(
    document_id: UUID,
    format: str = "pdf",
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user)
):
    """Generate a fresh signed URL for document download."""
    from app.models.document import Document
    
    document = db.query(Document).filter(
        Document.id == document_id,
        Document.organization_id == current_user.organization_id
    ).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    from app.core.supabase import supabase_client
    
    # Generate fresh signed URL (1 hour)
    if format == "pdf" and document.pdf_path:
        url_data = supabase_client.storage.from_("documents").create_signed_url(
            document.pdf_path, 3600
        )
        return {"url": url_data.get("signedURL")}
    elif format == "docx" and document.docx_path:
        url_data = supabase_client.storage.from_("documents").create_signed_url(
            document.docx_path, 3600
        )
        return {"url": url_data.get("signedURL")}
    else:
        raise HTTPException(status_code=400, detail="Invalid format or document not available")


@router.get("/types")
def list_document_types():
    """List available document types and their requirements."""
    return {
        doc_type: {
            "title": config["title"],
            "description": config["description"],
            "requires_use_case": config["requires_use_case"],
            "obligatory_for": config["obligatory_for"]
        }
        for doc_type, config in DOCUMENT_TYPES.items()
    }
