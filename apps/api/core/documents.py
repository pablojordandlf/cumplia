#!/usr/bin/env python3
"""
Document Generation Engine for CumplIA
Generates compliance documents in PDF and DOCX formats
"""

from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML, CSS
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import sys
import json
import argparse


class DocumentGenerator:
    """Generates compliance documents from templates"""
    
    def __init__(self, templates_dir='/tmp/cumplia/apps/api/templates'):
        self.env = Environment(
            loader=FileSystemLoader(templates_dir),
            autoescape=True
        )
        self.templates_dir = templates_dir
    
    def render_template(self, template_name, data):
        """Render HTML template with Jinja2"""
        template = self.env.get_template(f'{template_name}.html')
        return template.render(**data)
    
    def generate_pdf(self, template_name, data, output_path):
        """Generate PDF from HTML template"""
        html_content = self.render_template(template_name, data)
        HTML(string=html_content).write_pdf(output_path)
        return output_path
    
    def generate_docx(self, template_name, data, output_path):
        """Generate DOCX document"""
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add header with company name
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_run = header.add_run(data.get('company_name', 'Company Name'))
        company_run.bold = True
        company_run.font.size = Pt(18)
        company_run.font.color.rgb = RGBColor(45, 55, 72)
        
        # Add title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run('Artificial Intelligence Policy')
        title_run.bold = True
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor(26, 54, 93)
        
        # Add effective date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"Effective Date: {data.get('effective_date', 'N/A')}")
        date_run.font.size = Pt(11)
        date_run.font.color.rgb = RGBColor(113, 128, 150)
        
        doc.add_paragraph()  # Spacing
        
        # Section 1: Purpose and Scope
        self._add_section_heading(doc, '1. Purpose and Scope')
        self._add_paragraph(doc, data.get('purpose_scope', ''))
        
        # Section 2: AI System Inventory
        self._add_section_heading(doc, '2. AI System Inventory')
        self._add_paragraph(doc, 'This policy covers the following AI systems:')
        ai_systems = data.get('ai_systems', [])
        for system in ai_systems:
            self._add_bullet_point(doc, f"{system.get('name', '')} - {system.get('description', '')} (Risk Level: {system.get('risk_level', '')})")
        
        # Section 3: Compliance Framework
        self._add_section_heading(doc, '3. Compliance Framework')
        self._add_paragraph(doc, 'This policy ensures compliance with:')
        self._add_bullet_point(doc, 'EU AI Act Regulation (EU) 2024/1689')
        self._add_bullet_point(doc, 'GDPR where applicable')
        self._add_bullet_point(doc, 'Industry-specific regulations')
        
        # Section 4: Governance Structure
        self._add_section_heading(doc, '4. Governance Structure')
        self._add_paragraph(doc, data.get('governance_structure', ''))
        
        # Section 5: Risk Management
        self._add_section_heading(doc, '5. Risk Management')
        self._add_paragraph(doc, data.get('risk_management_approach', ''))
        
        # Section 6: Human Oversight
        self._add_section_heading(doc, '6. Human Oversight')
        self._add_paragraph(doc, data.get('human_oversight_measures', ''))
        
        # Section 7: Data Governance
        self._add_section_heading(doc, '7. Data Governance')
        self._add_paragraph(doc, data.get('data_governance', ''))
        
        # Section 8: Transparency and Documentation
        self._add_section_heading(doc, '8. Transparency and Documentation')
        self._add_paragraph(doc, data.get('transparency_measures', ''))
        
        # Section 9: Review and Updates
        self._add_section_heading(doc, '9. Review and Updates')
        self._add_paragraph(doc, f"This policy will be reviewed {data.get('review_frequency', 'annually')}.")
        
        # Add signature section
        doc.add_paragraph()
        doc.add_paragraph()
        sig_para = doc.add_paragraph()
        sig_para.add_run(f"Approved by: {data.get('approver_name', '')}").bold = True
        doc.add_paragraph(f"Date: {data.get('approval_date', '')}")
        
        doc.save(output_path)
        return output_path
    
    def _add_section_heading(self, doc, text):
        """Add a section heading"""
        heading = doc.add_heading(text, level=2)
        heading.runs[0].font.color.rgb = RGBColor(44, 82, 130)
        heading.runs[0].font.size = Pt(14)
    
    def _add_paragraph(self, doc, text):
        """Add a paragraph"""
        para = doc.add_paragraph(text)
        para.paragraph_format.line_spacing = 1.15
    
    def _add_bullet_point(self, doc, text):
        """Add a bullet point"""
        bullet = doc.add_paragraph(text, style='List Bullet')
        bullet.paragraph_format.left_indent = Inches(0.25)


def main():
    """CLI entry point for document generation"""
    parser = argparse.ArgumentParser(description='Generate compliance documents')
    parser.add_argument('command', choices=['generate'], help='Command to execute')
    parser.add_argument('template_type', help='Template type (e.g., ai_policy_en)')
    parser.add_argument('format', choices=['pdf', 'docx'], help='Output format')
    parser.add_argument('data', help='JSON data for template')
    parser.add_argument('output_path', help='Output file path')
    
    args = parser.parse_args()
    
    if args.command == 'generate':
        generator = DocumentGenerator()
        data = json.loads(args.data)
        
        if args.format == 'pdf':
            generator.generate_pdf(args.template_type, data, args.output_path)
        else:
            generator.generate_docx(args.template_type, data, args.output_path)
        
        print(f"Document generated: {args.output_path}")


if __name__ == '__main__':
    main()
